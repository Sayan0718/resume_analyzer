# Flask backend with job recommendation and ATS logic
from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from pdf_extractor import extract_text_from_pdf
from model_utils import recommend_jobs
from model_utils import calculate_ats_score
from docx import Document
import google.generativeai as genai
from config import gemini_api_key
import markdown
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt


app = Flask(__name__)
import os

UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


@app.route("/")
def home():
    return render_template("index.html")

@app.route("/job_recommendation", methods=["GET", "POST"])
def job_recommendation():
    if request.method == "POST":
        file = request.files["resume"]
        filename = secure_filename(file.filename)
        filepath = f"uploads/{filename}"
        file.save(filepath)

        resume_text = extract_text_from_pdf(filepath)
        jobs = recommend_jobs(resume_text)
        return render_template("results.html", jobs=jobs)

    return render_template("job_recommendation.html")

@app.route("/ats_score_checker", methods=["GET", "POST"])
def ats_score_checker():
    if request.method == "POST":
        file = request.files["resume"]
        job_description = request.form.get("job_description", "")


        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        resume_text = extract_text_from_pdf(filepath)

        # Calculate ATS Score (simple example based on skill matching)
        ats_score, missing_skills = calculate_ats_score(resume_text, job_description)

        return render_template("results.html", ats_score=ats_score, missing_skills=missing_skills)

    return render_template("ats_score_checker.html")


@app.route("/resume_checker", methods=["GET", "POST"])
def resume_checker():
    if request.method == "POST":
        file = request.files["resume"]
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(filepath)

        resume_text = extract_text_from_pdf(filepath)

        # For now we’ll do a basic improvement suggestion logic:
        from model_utils import suggest_improvements
        missing_skills = suggest_improvements(resume_text)

        return render_template("results.html", missing_skills=missing_skills)

    return render_template("resume_checker.html")

@app.route("/build_resume")
def build_resume():
    return render_template("resume_builder.html")

@app.route("/generate_resume", methods=["POST"])
def generate_resume():
    data = request.form
    name = data['name']
    email = data['email']
    phone = data['phone']
    linkedin = data['linkedin']
    github = data['github']
    objective = data['objective']
    skills = data['skills'].split(",")
    education = data['education']
    projects = data['projects'].split("\n")
    certifications = data['certifications'].split("\n")

    # optional
    achievements = data.get('achievements', '').split("\n") if 'achievements' in data else []
    conferences = data.get('conferences', '').split("\n") if 'conferences' in data else []

    doc = Document()
    doc.add_heading(name, 0)

    p = doc.add_paragraph()
    p.add_run(f"Email: {email} | Phone: {phone}\n").bold = True
    p.add_run(f"LinkedIn: {linkedin} | GitHub: {github}")

    doc.add_heading('OBJECTIVE', level=1)
    doc.add_paragraph(objective)

    doc.add_heading('TECHNICAL SKILLS', level=1)
    for skill in skills:
        if skill.strip():
            doc.add_paragraph(skill.strip(), style='ListBullet')

    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph(education)

    if projects and any(projects):
        doc.add_heading('PROJECTS', level=1)
        for proj in projects:
            if proj.strip():
                doc.add_paragraph(proj.strip(), style='ListBullet')

    if certifications and any(certifications):
        doc.add_heading('CERTIFICATIONS', level=1)
        for cert in certifications:
            if cert.strip():
                doc.add_paragraph(cert.strip(), style='ListBullet')

    if achievements and any(achievements):
        doc.add_heading('ACHIEVEMENTS', level=1)
        for ach in achievements:
            if ach.strip():
                doc.add_paragraph(ach.strip(), style='ListBullet')

    if conferences and any(conferences):
        doc.add_heading('CONFERENCES', level=1)
        for conf in conferences:
            if conf.strip():
                doc.add_paragraph(conf.strip(), style='ListBullet')

    # formatting
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    output_path = f"uploads/{secure_filename(name)}_Resume.docx"
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)

genai.configure(
    api_key=gemini_api_key,
    transport='rest',    # Force REST instead of default GRPC
)
@app.route('/mock_interview', methods=['GET', 'POST'])
def mock_interview():
    if request.method == 'POST':
        role = request.form.get('role')
        technologies = request.form.get('technologies')

        prompt = f"Generate 5 interview questions along with brief answers for a {role} role focusing on {technologies}. Format the response cleanly in markdown."

        try:
            model = genai.GenerativeModel('models/gemini-1.5-pro')
            response = model.generate_content(prompt)

            # ✅ Convert markdown response to HTML before rendering
            generated_text = response.text
            html_output = markdown.markdown(generated_text)

            return render_template('mock_interview_result.html', questions=html_output)

        except Exception as e:
            return f"❌ Error generating interview using Gemini: {e}"

    return render_template('mock_interview.html')

if __name__ == "__main__":
    app.run(debug=True)
